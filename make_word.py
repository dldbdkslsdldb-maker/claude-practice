import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 기본 여백 설정 ────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

# ── 헬퍼 함수 ─────────────────────────────────────────────
def set_font(run, bold=False, size=11, color=None, font_name='맑은 고딕'):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_font(run, bold=True, size=16, color=(31, 73, 125))
    elif level == 2:
        set_font(run, bold=True, size=13, color=(47, 84, 150))
    elif level == 3:
        set_font(run, bold=True, size=11, color=(68, 114, 196))
    # 단락 간격
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=10.5)
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=10.5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.8)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B0B0B0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)

def add_image(doc, path, width_cm=15):
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Cm(width_cm))
        p.paragraph_format.space_after = Pt(6)
    except Exception as e:
        add_body(doc, f'[이미지 삽입 실패: {e}]')

def shade_row(row, hex_color='D6E4F7'):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def cell_text(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_font(run, bold=bold, size=size)


# ══════════════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════════════
doc.add_paragraph('\n\n')
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('데이터 분석 실습 보고서')
set_font(r, bold=True, size=22, color=(31, 73, 125))

doc.add_paragraph()
t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('챕터 3 숙제 | 유창현')
set_font(r2, bold=False, size=14, color=(89, 89, 89))

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('분석 파일: 챕터3숙제_데이터분석.xlsx')
set_font(r3, size=11, color=(127, 127, 127))

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 1. 데이터 구조 파악
# ══════════════════════════════════════════════════════════
add_heading(doc, '1. 데이터 구조 파악', level=1)
add_divider(doc)

add_heading(doc, '1.1 기본 정보', level=2)
tbl = doc.add_table(rows=5, cols=2)
tbl.style = 'Table Grid'
set_col_widths(tbl, [4.5, 10.5])
shade_row(tbl.rows[0], 'D6E4F7')
headers = [('항목', '내용'),
           ('시트명', '월별판매데이터'),
           ('총 행 수', '220건 (헤더 제외)'),
           ('총 열 수', '9개'),
           ('분석 기간', '2026년 1월 1일 ~ 2026년 4월 30일')]
for i, (k, v) in enumerate(headers):
    cell_text(tbl.rows[i].cells[0], k, bold=(i==0), size=10)
    cell_text(tbl.rows[i].cells[1], v, bold=(i==0), size=10)

doc.add_paragraph()
add_heading(doc, '1.2 컬럼 정보', level=2)
cols_info = [
    ('컬럼명',    '데이터 유형', '결측값', '설명'),
    ('주문일자',  'datetime',   '0개',   '주문 날짜'),
    ('주문번호',  'str',        '0개',   '고유 주문 ID (예: ORD-20260004)'),
    ('지역',      'str',        '0개',   '주문 발생 지역 (7개 시/도)'),
    ('카테고리',  'str',        '0개',   '상품 분류 (4종)'),
    ('상품명',    'str',        '0개',   '판매 상품명 (15종)'),
    ('수량',      'int',        '0개',   '주문 수량 (1~7개)'),
    ('단가',      'int',        '0개',   '상품 단가 (원)'),
    ('매출',      'str/수식 ⚠️','0개',   '=수량×단가 수식 — 직접 재계산 필요'),
    ('판매채널',  'str',        '0개',   '판매 경로 (5개 채널)'),
]
tbl2 = doc.add_table(rows=len(cols_info), cols=4)
tbl2.style = 'Table Grid'
set_col_widths(tbl2, [3.0, 3.0, 2.0, 7.0])
shade_row(tbl2.rows[0], 'D6E4F7')
for i, row_data in enumerate(cols_info):
    for j, val in enumerate(row_data):
        cell_text(tbl2.rows[i].cells[j], val, bold=(i==0), size=10)

doc.add_paragraph()
add_heading(doc, '1.3 데이터 품질 이슈', level=2)
add_bullet(doc, '⚠️ 매출 컬럼이 =F2*G2 형태의 Excel 수식 문자열로 저장되어 있어 openpyxl로 읽으면 숫자가 아닌 문자열 반환')
add_bullet(doc, '처리 방법: 수량(F열) × 단가(G열)를 Python에서 직접 곱하여 매출액 재계산 후 분석에 사용')
add_bullet(doc, '결측값 없음 — 전 컬럼 220건 완전 수집')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 2. 데이터 분석 결과
# ══════════════════════════════════════════════════════════
add_heading(doc, '2. 데이터 분석 결과', level=1)
add_divider(doc)

add_heading(doc, '2.1 주요 KPI', level=2)
kpi_data = [
    ('지표',          '값'),
    ('전체 총매출',    '80.5백만원'),
    ('평균 단가',      '164,091원'),
    ('평균 주문 수량', '2.25개'),
    ('최고 단가 상품', '459,000원'),
    ('최저 단가 상품', '29,000원'),
    ('고유 상품 수',   '15개'),
]
tbl3 = doc.add_table(rows=len(kpi_data), cols=2)
tbl3.style = 'Table Grid'
set_col_widths(tbl3, [5.5, 9.5])
shade_row(tbl3.rows[0], 'D6E4F7')
for i, (k, v) in enumerate(kpi_data):
    cell_text(tbl3.rows[i].cells[0], k, bold=(i==0), size=10)
    cell_text(tbl3.rows[i].cells[1], v, bold=(i==0), size=10)

doc.add_paragraph()
add_heading(doc, '2.2 월별 매출 집계', level=2)
month_data = [
    ('월',  '총매출',    '주문건수', '전월 대비'),
    ('1월', '18.0M원',  '62건',    '-'),
    ('2월', '16.4M원',  '43건',    '▼ -8.9%'),
    ('3월', '26.3M원',  '65건',    '▲ +60.4%'),
    ('4월', '19.7M원',  '50건',    '▼ -25.1%'),
]
tbl4 = doc.add_table(rows=len(month_data), cols=4)
tbl4.style = 'Table Grid'
set_col_widths(tbl4, [2.5, 3.5, 3.0, 6.0])
shade_row(tbl4.rows[0], 'D6E4F7')
for i, row_data in enumerate(month_data):
    for j, val in enumerate(row_data):
        cell_text(tbl4.rows[i].cells[j], val, bold=(i==0), size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER if j in [0,2] else WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph()
add_heading(doc, '2.3 카테고리별 집계', level=2)
cat_data = [
    ('카테고리',  '주문건수', '비율'),
    ('전자기기',  '84건',    '38.2%'),
    ('생활가전',  '68건',    '30.9%'),
    ('주방용품',  '38건',    '17.3%'),
    ('뷰티',      '30건',    '13.6%'),
]
tbl5 = doc.add_table(rows=len(cat_data), cols=3)
tbl5.style = 'Table Grid'
set_col_widths(tbl5, [4.0, 4.0, 4.0])
shade_row(tbl5.rows[0], 'D6E4F7')
for i, row_data in enumerate(cat_data):
    for j, val in enumerate(row_data):
        cell_text(tbl5.rows[i].cells[j], val, bold=(i==0), size=10)

doc.add_paragraph()
add_heading(doc, '2.4 판매채널별 집계', level=2)
ch_data = [
    ('판매채널',          '매출',     '비중'),
    ('쿠팡',              '31.7M원',  '39.4%'),
    ('네이버스마트스토어', '22.5M원',  '28.0%'),
    ('자사몰',            '18.4M원',  '22.8%'),
    ('11번가',            '4.6M원',   '5.8%'),
    ('오프라인매장',       '3.3M원',   '4.1%'),
]
tbl6 = doc.add_table(rows=len(ch_data), cols=3)
tbl6.style = 'Table Grid'
set_col_widths(tbl6, [5.0, 4.0, 4.0])
shade_row(tbl6.rows[0], 'D6E4F7')
for i, row_data in enumerate(ch_data):
    for j, val in enumerate(row_data):
        cell_text(tbl6.rows[i].cells[j], val, bold=(i==0), size=10)

doc.add_paragraph()
add_heading(doc, '2.5 지역별 집계', level=2)
reg_data = [
    ('지역', '주문건수', '비율'),
    ('서울', '76건', '34.5%'),
    ('경기', '63건', '28.6%'),
    ('부산', '27건', '12.3%'),
    ('대구', '17건', '7.7%'),
    ('인천', '15건', '6.8%'),
    ('광주', '15건', '6.8%'),
    ('대전', '7건',  '3.2%'),
]
tbl7 = doc.add_table(rows=len(reg_data), cols=3)
tbl7.style = 'Table Grid'
set_col_widths(tbl7, [4.0, 4.0, 4.0])
shade_row(tbl7.rows[0], 'D6E4F7')
for i, row_data in enumerate(reg_data):
    for j, val in enumerate(row_data):
        cell_text(tbl7.rows[i].cells[j], val, bold=(i==0), size=10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 3. 생성된 차트
# ══════════════════════════════════════════════════════════
add_heading(doc, '3. 생성된 차트', level=1)
add_divider(doc)

BASE = r'c:\Users\유창현\Desktop\claude-practice'

add_heading(doc, '3.1 선 차트 — 월별 카테고리별 매출 추이', level=2)
add_body(doc, '날짜 컬럼(주문일자) + 수치 컬럼(매출액)이 존재하여 시간적 흐름을 파악하기 위해 선 차트를 선택하였다.')
add_image(doc, f'{BASE}/chart_line.png', width_cm=14)
add_bullet(doc, '3월에 전체 합계 26.3M원으로 최고점 도달 — 전자기기가 2월(3.2M) → 3월(9.2M)으로 3배 급등')
add_bullet(doc, '생활가전은 3~4월 11M대를 꾸준히 유지하며 가장 안정적인 성장세')
add_bullet(doc, '주방용품·뷰티는 3월 이후 급락하여 계절성 또는 프로모션 종료 영향으로 추정')

doc.add_paragraph()
add_heading(doc, '3.2 누적 막대 차트 — 월별 카테고리별 매출 구성', level=2)
add_body(doc, '범주형 2개 이상(월, 카테고리) + 수치 컬럼이 있어 각 월의 카테고리 구성 비중을 한눈에 보기 위해 누적 막대 차트를 선택하였다.')
add_image(doc, f'{BASE}/chart_bar.png', width_cm=14)
add_bullet(doc, '2월은 모든 카테고리가 동반 하락 — 전월 대비 전체 -9% 감소')
add_bullet(doc, '3월 막대에서 전자기기·생활가전이 전체의 79% 차지 → 특정 카테고리 쏠림 심화')
add_bullet(doc, '주방용품(초록)은 1~2월 대비 3~4월에 눈에 띄게 얇아져 수요 감소 확인 가능')

doc.add_paragraph()
add_heading(doc, '3.3 원형 차트 — 판매채널별 매출 비중', level=2)
add_body(doc, '판매채널이라는 단일 범주형 컬럼의 비중 구조를 파악하기 위해 원형 차트를 선택하였다.')
add_image(doc, f'{BASE}/chart_pie.png', width_cm=12)
add_bullet(doc, '쿠팡이 39.4%(31.7M)로 압도적 1위 — 단일 채널 의존도가 높아 리스크 요인')
add_bullet(doc, '상위 3개 채널(쿠팡+네이버+자사몰)이 전체의 90% 이상 점유')
add_bullet(doc, '오프라인매장은 4.1%(3.3M)에 불과 — 사실상 온라인 전용 사업 구조')

doc.add_paragraph()
add_heading(doc, '3.4 수평 막대 차트 — 지역별 주문 건수', level=2)
add_body(doc, '지역이라는 범주형 컬럼의 건수 순위를 직관적으로 보기 위해 수평 막대 차트를 선택하였다.')
add_image(doc, f'{BASE}/chart_hbar.png', width_cm=14)
add_bullet(doc, '서울(76건)+경기(63건)=139건으로 전체의 63% — 수도권 집중도 극심')
add_bullet(doc, '부산(27건)이 3위이나 서울 대비 35% 수준에 그침')
add_bullet(doc, '대전(7건)은 최하위로 지방 광역시 중 유일한 한 자릿수 — 마케팅 공백 지역')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 4. 작성한 Skill 내용
# ══════════════════════════════════════════════════════════
add_heading(doc, '4. 작성한 Skill 내용', level=1)
add_divider(doc)

add_heading(doc, '4.1 Skill 기본 정보', level=2)
skill_info = [
    ('항목',      '내용'),
    ('Skill 이름', 'data-analysis'),
    ('파일 경로',  '.claude/skills/data anlysis/SKILL.md'),
    ('설명',       'Excel/CSV 데이터를 받아 구조 파악 → 차트 생성 → 인사이트 출력을 순서대로 자동 수행'),
    ('기술 스택',  'Python + openpyxl + matplotlib'),
]
tbl8 = doc.add_table(rows=len(skill_info), cols=2)
tbl8.style = 'Table Grid'
set_col_widths(tbl8, [4.0, 11.0])
shade_row(tbl8.rows[0], 'D6E4F7')
for i, (k, v) in enumerate(skill_info):
    cell_text(tbl8.rows[i].cells[0], k, bold=(i==0), size=10)
    cell_text(tbl8.rows[i].cells[1], v, bold=(i==0), size=10)

doc.add_paragraph()
add_heading(doc, '4.2 Skill 실행 절차 (5단계)', level=2)
steps = [
    ('1단계', '데이터 구조 파악',
     '시트명, 행×열, 컬럼명·유형·결측값 확인. 수식 셀(=F*G) 자동 감지 후 ⚠️ 표시.'),
    ('2단계', '집계 분석',
     '날짜→기간/월별, 범주형→건수/비율, 수치형→최소/최대/평균, 금액 컬럼→범주별 합계 자동 산출.'),
    ('3단계', '차트 유형 자동 선택 및 생성',
     '날짜+수치→선 차트, 범주형≥2+수치→누적 막대, 범주 비중→원형, 건수 순위→수평 막대. 복수 패턴이면 모두 생성.'),
    ('4단계', '인사이트 출력',
     '데이터 개요 / 주요 수치 / 차트별 발견사항 / 주목할 패턴 / 데이터 품질 이슈 순으로 대화창 출력.'),
    ('5단계', '완료 보고',
     '✅ 구조 파악 / ✅ 차트 생성 / ✅ 인사이트 출력 완료 체크리스트 출력.'),
]
tbl9 = doc.add_table(rows=len(steps)+1, cols=3)
tbl9.style = 'Table Grid'
set_col_widths(tbl9, [1.8, 3.5, 9.7])
shade_row(tbl9.rows[0], 'D6E4F7')
cell_text(tbl9.rows[0].cells[0], '단계',    bold=True, size=10)
cell_text(tbl9.rows[0].cells[1], '이름',    bold=True, size=10)
cell_text(tbl9.rows[0].cells[2], '수행 내용', bold=True, size=10)
for i, (step, name, desc) in enumerate(steps):
    cell_text(tbl9.rows[i+1].cells[0], step, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(tbl9.rows[i+1].cells[1], name, bold=True, size=10)
    cell_text(tbl9.rows[i+1].cells[2], desc, size=10)

doc.add_paragraph()
add_heading(doc, '4.3 예외 처리 규칙', level=2)
exc_data = [
    ('상황',                   '처리 방법'),
    ('openpyxl 미설치',         'pip install openpyxl 실행 후 재시도'),
    ('matplotlib 미설치',       'pip install matplotlib 실행 후 재시도'),
    ('수식 셀(str =...)',       '수량×단가 등 인접 수치 컬럼으로 직접 계산'),
    ('날짜 컬럼 없음',           '월별 추이 차트 생략, 나머지 차트만 생성'),
    ('수치 컬럼 없음',           '건수 기반 막대/원형 차트만 생성'),
    ('한글 폰트 없음',           '영문 폰트로 대체 후 "한글 폰트 미탐지" 안내'),
    ('범주 고유값 > 20개',       '상위 10개만 표시, 나머지 "기타"로 묶음'),
]
tbl10 = doc.add_table(rows=len(exc_data), cols=2)
tbl10.style = 'Table Grid'
set_col_widths(tbl10, [5.5, 9.5])
shade_row(tbl10.rows[0], 'D6E4F7')
for i, (k, v) in enumerate(exc_data):
    cell_text(tbl10.rows[i].cells[0], k, bold=(i==0), size=10)
    cell_text(tbl10.rows[i].cells[1], v, bold=(i==0), size=10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 5. Skill 실행 결과
# ══════════════════════════════════════════════════════════
add_heading(doc, '5. Skill 실행 결과', level=1)
add_divider(doc)

add_heading(doc, '5.1 실행 환경', level=2)
env_data = [
    ('항목',      '내용'),
    ('입력 파일', '챕터3숙제_데이터분석.xlsx'),
    ('Python',    '3.x + openpyxl + matplotlib 3.11.0'),
    ('한글 폰트', 'Malgun Gothic (자동 탐지 성공)'),
    ('저장 경로', 'c:\\Users\\유창현\\Desktop\\claude-practice\\'),
]
tbl11 = doc.add_table(rows=len(env_data), cols=2)
tbl11.style = 'Table Grid'
set_col_widths(tbl11, [4.0, 11.0])
shade_row(tbl11.rows[0], 'D6E4F7')
for i, (k, v) in enumerate(env_data):
    cell_text(tbl11.rows[i].cells[0], k, bold=(i==0), size=10)
    cell_text(tbl11.rows[i].cells[1], v, bold=(i==0), size=10)

doc.add_paragraph()
add_heading(doc, '5.2 단계별 실행 로그', level=2)
add_body(doc, '아래는 Skill 실행 중 출력된 주요 로그 메시지이다.')

log_lines = [
    '[폰트] Malgun Gothic 적용',
    '[구조] 220건 × 9개 컬럼 로드 완료',
    '[수식] 매출 컬럼 수식 감지 → 수량×단가로 재계산 완료',
    'chart_line.png 저장 완료',
    'chart_bar.png 저장 완료',
    'chart_pie.png 저장 완료',
    'chart_hbar.png 저장 완료',
    '월별 합계(M): {1월: 18.0, 2월: 16.4, 3월: 26.3, 4월: 19.7}',
    '채널별 매출(M): {쿠팡: 31.7, 자사몰: 18.4, 네이버스마트스토어: 22.5, 11번가: 4.6, 오프라인매장: 3.3}',
    '전체 총매출: 80.5M원',
    '평균 단가: 164,091원',
]
p_log = doc.add_paragraph()
p_log.paragraph_format.left_indent = Cm(0.8)
p_log.paragraph_format.space_after = Pt(4)
for line in log_lines:
    run = p_log.add_run(line + '\n')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(40, 40, 40)

doc.add_paragraph()
add_heading(doc, '5.3 완료 보고', level=2)
result_lines = [
    '✅ 구조 파악: 220건 × 9개 컬럼',
    '✅ 차트 생성: chart_line.png / chart_bar.png / chart_pie.png / chart_hbar.png',
    '✅ 인사이트 출력 완료',
]
for line in result_lines:
    add_bullet(doc, line)

doc.add_paragraph()
add_heading(doc, '5.4 주요 인사이트 요약', level=2)
insights = [
    '3월이 전체 고점(26.3M) — 전자기기가 2월 대비 3배 급등하여 전체 매출 견인',
    '2월 비수기 패턴 뚜렷 — 전월 대비 -9% 감소, 건수도 62건→43건으로 31% 감소',
    '쿠팡 단일 채널 의존도 39.4% — 채널 다각화 전략 검토 필요',
    '수도권(서울+경기) 집중도 63% — 지방 광역시 중 대전(7건) 마케팅 공백 존재',
    '주방용품·뷰티 3월 이후 급락 — 계절성 또는 프로모션 종료 영향으로 추정',
    '뷰티 4월 매출 75% 급감(4.5M → 1.1M) — 원인 파악 및 4월 판촉 전략 재검토 필요',
]
for ins in insights:
    add_bullet(doc, ins)

# ── 저장 ──────────────────────────────────────────────────
OUT = r'c:\Users\유창현\Desktop\claude-practice\챕터3숙제_유창현.docx'
doc.save(OUT)
print(f'저장 완료: {OUT}')
